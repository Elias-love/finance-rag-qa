"""模拟数据生成器：生成虚构"星辰集团"的财务报表Excel + 制度文档docx到 data/uploads/

用法：
    python scripts/generate_mock_data.py      # 生成模拟文件
    python rebuild_index.py                   # 走真实解析管道入库（SQLite + ChromaDB）

数据特点：
- 全部虚构，与任何真实企业无关；数字自洽（资产=负债+权益、净利润=利润总额-所得税）
- 覆盖系统全部演示场景：单体/合并双口径、海外公司中英文桥接、双币种报表
- 确定性生成（无随机数），eval/golden_set.jsonl 的期望值与之精确对应
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

import openpyxl
from docx import Document

from config import UPLOAD_DIR

# ============================================================
# 虚构公司体系：收入基数 + 成本费用率（确定性，改动会导致评估集失配）
# ============================================================
COMPANIES = {
    # 文件名: (营业收入, 成本率, 销售费用率, 管理费用率, 研发费用率, 所得税率)
    "01深圳星辰数字科技集团股份有限公司_2025TB.xlsx": (580_231_456.78, 0.62, 0.055, 0.085, 0.048, 0.15),
    "合并0：星辰集团_2025TB.xlsx":                  (1_352_648_913.24, 0.60, 0.060, 0.080, 0.052, 0.15),
    "02深圳市辰拓智能设备有限公司_2025TB.xlsx":      (321_574_208.65, 0.65, 0.050, 0.070, 0.040, 0.15),
    "合并1：辰拓合并_2025TB.xlsx":                  (392_837_461.32, 0.64, 0.052, 0.072, 0.042, 0.15),
    "03深圳星美智能材料有限公司_2025TB.xlsx":        (263_918_570.41, 0.58, 0.065, 0.075, 0.055, 0.15),
    "05深圳星锐精密装备有限公司_2025TB.xlsx":        (152_406_385.19, 0.67, 0.045, 0.090, 0.038, 0.25),
    "08STARNOVA (HONG KONG) LIMITED_2025TB.xlsx":   (96_235_147.83, 0.72, 0.030, 0.050, 0.000, 0.165),
    "08-1STARNOVA MALAYSIA SDN. BHD._2025TB（人民币）.xlsx": (47_582_931.06, 0.70, 0.040, 0.060, 0.000, 0.24),
    "08-1STARNOVA MALAYSIA SDN. BHD._2025TB（林吉特）.xlsx": (30_698_665.20, 0.70, 0.040, 0.060, 0.000, 0.24),
    "11惠州星辰实业有限公司_2025TB.xlsx":            (213_754_692.37, 0.75, 0.035, 0.045, 0.025, 0.25),
}

R2 = lambda x: round(x, 2)
PRIOR = 0.88  # 上年同期数 = 本期 × 0.88（简化的同比基数）


def income_statement(revenue, costr, sf, mf, rd, taxr):
    """利润表：从收入基数推导全部行项目，保证钩稽自洽"""
    rows = {}
    rows["一、营业收入"] = R2(revenue)
    rows["    减：营业成本"] = R2(revenue * costr)
    rows["        税金及附加"] = R2(revenue * 0.008)
    rows["        销售费用"] = R2(revenue * sf)
    rows["        管理费用"] = R2(revenue * mf)
    rows["        研发费用"] = R2(revenue * rd)
    rows["        财务费用"] = R2(revenue * 0.006)
    op = rows["一、营业收入"] - sum(
        rows[k] for k in list(rows)[1:])
    rows["二、营业利润"] = R2(op)
    rows["    加：营业外收入"] = R2(revenue * 0.001)
    rows["    减：营业外支出"] = R2(revenue * 0.0005)
    total = rows["二、营业利润"] + rows["    加：营业外收入"] - rows["    减：营业外支出"]
    rows["三、利润总额"] = R2(total)
    rows["    减：所得税费用"] = R2(total * taxr)
    rows["四、净利润"] = R2(rows["三、利润总额"] - rows["    减：所得税费用"])
    return rows


def balance_sheet(revenue, net_profit):
    """资产负债表：左右两栏，资产总计 = 负债和所有者权益总计"""
    assets = {}
    assets["货币资金"] = R2(revenue * 0.22)
    assets["应收账款"] = R2(revenue * 0.31)
    assets["存货"] = R2(revenue * 0.26)
    assets["其他流动资产"] = R2(revenue * 0.04)
    assets["流动资产合计"] = R2(sum(assets.values()))
    assets["固定资产"] = R2(revenue * 0.35)
    assets["无形资产"] = R2(revenue * 0.08)
    assets["资产总计"] = R2(assets["流动资产合计"] + assets["固定资产"] + assets["无形资产"])

    liab = {}
    liab["短期借款"] = R2(revenue * 0.12)
    liab["应付账款"] = R2(revenue * 0.24)
    liab["其他应付款"] = R2(revenue * 0.05)
    liab["流动负债合计"] = R2(sum(liab.values()))
    liab["长期借款"] = R2(revenue * 0.10)
    liab["负债合计"] = R2(liab["流动负债合计"] + liab["长期借款"])
    liab["实收资本"] = R2(revenue * 0.30)
    liab["资本公积"] = R2(revenue * 0.10)
    # 未分配利润为轧差项，强制资产=负债+权益
    liab["未分配利润"] = R2(assets["资产总计"] - liab["负债合计"]
                        - liab["实收资本"] - liab["资本公积"])
    liab["所有者权益合计"] = R2(liab["实收资本"] + liab["资本公积"] + liab["未分配利润"])
    liab["负债和所有者权益总计"] = R2(liab["负债合计"] + liab["所有者权益合计"])
    return assets, liab


def cash_flow(revenue, net_profit):
    """现金流量表：三大活动净额，简化口径"""
    rows = {}
    rows["经营活动产生的现金流量净额"] = R2(net_profit * 1.12)
    rows["投资活动产生的现金流量净额"] = R2(-revenue * 0.06)
    rows["筹资活动产生的现金流量净额"] = R2(revenue * 0.02)
    rows["现金及现金等价物净增加额"] = R2(sum(rows.values()))
    return rows


def write_workbook(path: Path, company_name: str, params):
    revenue, costr, sf, mf, rd, taxr = params
    inc = income_statement(revenue, costr, sf, mf, rd, taxr)
    assets, liab = balance_sheet(revenue, inc["四、净利润"])
    cf = cash_flow(revenue, inc["四、净利润"])

    wb = openpyxl.Workbook()

    # —— 利润表 ——
    ws = wb.active
    ws.title = "利润表"
    ws.append(["利润表"])
    ws.append([f"编制单位：{company_name}", None, "单位：元"])
    ws.append(["项目", "本期数", "上年同期数"])
    for item, val in inc.items():
        ws.append([item, val, R2(val * PRIOR)])

    # —— 资产负债表（左右两栏）——
    ws = wb.create_sheet("资产负债表")
    ws.append(["资产负债表"])
    ws.append([f"编制单位：{company_name}", None, None, None, None, "单位：元"])
    ws.append(["资产", "期末数", "上年年末数", "负债和所有者权益", "期末数", "上年年末数"])
    a_items, l_items = list(assets.items()), list(liab.items())
    for i in range(max(len(a_items), len(l_items))):
        a_name, a_val = a_items[i] if i < len(a_items) else ("", None)
        l_name, l_val = l_items[i] if i < len(l_items) else ("", None)
        ws.append([a_name, a_val, R2(a_val * PRIOR) if a_val else None,
                   l_name, l_val, R2(l_val * PRIOR) if l_val is not None else None])

    # —— 现金流量表 ——
    ws = wb.create_sheet("现金流量表")
    ws.append(["现金流量表"])
    ws.append([f"编制单位：{company_name}", None, "单位：元"])
    ws.append(["项目", "本期数", "上年同期数"])
    for item, val in cf.items():
        ws.append([item, val, R2(val * PRIOR)])

    wb.save(path)
    return inc, assets


def write_policy_docs():
    """生成制度类模拟文档（RAG文本通道演示语料）"""
    docs = {
        "财务部考勤制度.docx": [
            ("财务部考勤管理制度", 0),
            ("一、工作时间", 1),
            ("工作时间为周一至周五，早上8点30分至下午18点，中午休息时间为12:15-14:00。", None),
            ("二、迟到早退", 1),
            ("单次迟到或早退不得超过15分钟，且1个月内累计迟到、早退不得超过3次，超出按旷工半天处理。", None),
            ("三、调休", 1),
            ("加班可申请调休，调休的最小单位时长为0.5小时，需提前一天在OA系统提交申请并经部门负责人审批。", None),
            ("四、请假", 1),
            ("事假需提前一天申请；病假需在当天上午10点前告知直属上级，事后补交病假条。", None),
        ],
        "应付账款管理制度.docx": [
            ("应付账款管理制度", 0),
            ("一、对账", 1),
            ("每月结账前，应付会计在ERP系统的入库查询模块（作业代码AP-130）导出当月外购入库数据，账期日期为上月26日至本月25日，与供应商对账单逐笔核对。", None),
            ("所有对账单核对无误后加盖财务专用章；如有扣款需另附扣款单。", None),
            ("二、发票审核", 1),
            ("收到供应商发票后，在发票审核模块（作业代码AP-110）按供应商检索，核对发票与入库登记信息一致后提交审核。", None),
            ("为保证职责分离，应付会计本人录入的发票不得由本人审核，须由总账会计复核。", None),
            ("三、暂估入账", 1),
            ("月末库存关账且当月需入账的进项发票全部审核完成后，方可执行暂估立账作业；当月已入库但未收到发票的入库单一律暂估入账。", None),
            ("四、月结", 1),
            ("应付模块月结前，须检查是否存在未审核、未过账单据，并核对应付模块与总账的勾稽关系；月结后如需修改数据，须执行反月结（月结还原）后方可调整。", None),
            ("五、付款", 1),
            ("付款给国外供应商时，付款信息上的汇率必须与原请款单汇率一致，出现差异须重新走请款审批。", None),
        ],
        "费用报销管理制度.docx": [
            ("费用报销管理制度", 0),
            ("一、报销时限", 1),
            ("费用发生后应在3个月内完成报销，跨年度费用原则上不予报销。", None),
            ("二、发票要求", 1),
            ("报销须取得合规发票，发票抬头须为公司全称，增值税专用发票须确保清晰可认证。", None),
            ("三、审批流程", 1),
            ("单笔5000元（含）以下由部门负责人审批；超过5000元须加签财务负责人；超过50000元须总经理审批。", None),
            ("四、差旅标准", 1),
            ("国内差旅住宿标准：一线城市每晚不超过450元，其他城市每晚不超过350元；市内交通费按实报销。", None),
        ],
    }
    for fname, blocks in docs.items():
        doc = Document()
        for text, level in blocks:
            if level == 0:
                doc.add_heading(text, level=0)
            elif level == 1:
                doc.add_heading(text, level=1)
            else:
                doc.add_paragraph(text)
        doc.save(UPLOAD_DIR / fname)
        print(f"  已生成 {fname}")


def main():
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    print("=" * 60)
    print("【生成模拟财务数据】（全部虚构，与任何真实企业无关）")
    print("=" * 60)

    print("\n① 财务报表 Excel：")
    key_figures = []
    for fname, params in COMPANIES.items():
        company = fname.split("_")[0].lstrip("0123456789-").lstrip("：")
        inc, assets = write_workbook(UPLOAD_DIR / fname, company, params)
        key_figures.append((fname, inc["一、营业收入"], inc["四、净利润"], assets["资产总计"]))
        print(f"  已生成 {fname}")

    print("\n② 制度文档 docx：")
    write_policy_docs()

    print("\n③ 关键数字（与 eval/golden_set.jsonl 对应）：")
    print(f"  {'文件':<52} {'营业收入':>16} {'净利润':>16} {'资产总计':>16}")
    for fname, rev, np_, ta in key_figures:
        print(f"  {fname:<52} {rev:>16,.2f} {np_:>16,.2f} {ta:>16,.2f}")

    print("\n完成。下一步执行: python rebuild_index.py")


if __name__ == "__main__":
    main()
