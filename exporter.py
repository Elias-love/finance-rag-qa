"""输出层：查询结果导出 Excel/CSV/PDF/Word（跨平台中文无乱码）"""

import re
from datetime import datetime
from pathlib import Path

import pandas as pd
from loguru import logger

from config import EXPORT_DIR, BASE_DIR, UPLOAD_DIR

# 内嵌中文字体路径（项目自带，macOS/Windows通用）
FONT_PATH = BASE_DIR / "data" / "NotoSansSC-Regular.ttf"


class Exporter:
    """查询结果导出：xlsx / csv / pdf / docx"""

    SUPPORTED_FORMATS = {"xlsx", "csv", "pdf", "docx"}

    def export(self, df: pd.DataFrame, filename: str = "", fmt: str = "xlsx") -> Path:
        if df is None or df.empty:
            raise ValueError("没有可导出的数据")

        if not filename:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"查询结果_{ts}"

        fmt = fmt.lower().strip(".")
        if fmt not in self.SUPPORTED_FORMATS:
            raise ValueError(f"不支持的导出格式: {fmt}，支持: {', '.join(sorted(self.SUPPORTED_FORMATS))}")

        path = EXPORT_DIR / f"{filename}.{fmt}"

        if fmt == "csv":
            self._export_csv(df, path)
        elif fmt == "xlsx":
            self._export_excel(df, path)
        elif fmt == "pdf":
            self._export_pdf(df, path)
        elif fmt == "docx":
            self._export_docx(df, path)

        logger.info(f"导出完成: {path} ({len(df)}行)")
        return path

    # ---------- 单sheet导出（保留原始格式，供来源下载）----------
    def export_single_sheet(self, source_file: str, sheet_name: str) -> Path:
        """从原始Excel中只导出命中的那一个sheet，尽量保留原格式"""
        src_path = UPLOAD_DIR / source_file
        if not src_path.exists():
            raise FileNotFoundError(f"原始文件不存在: {source_file}")

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_sheet = re.sub(r"[^\w一-鿿]", "_", str(sheet_name)) or "sheet"
        out_path = EXPORT_DIR / f"{src_path.stem}_{safe_sheet}_{ts}.xlsx"

        suffix = src_path.suffix.lower()

        if suffix in (".xlsx", ".xlsm"):
            # data_only=True 读取公式的缓存计算值，避免跨sheet公式删除后变 #REF!
            from openpyxl import load_workbook, Workbook
            src_wb = load_workbook(str(src_path), data_only=True)
            if sheet_name not in src_wb.sheetnames:
                raise ValueError(f"未找到工作表: {sheet_name}（可用: {src_wb.sheetnames}）")
            src_ws = src_wb[sheet_name]

            new_wb = Workbook()
            new_ws = new_wb.active
            new_ws.title = str(sheet_name)[:31]

            # 复制单元格的计算值
            none_count = 0
            total_count = 0
            for row in src_ws.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        new_ws[cell.coordinate] = cell.value
                        total_count += 1
                    else:
                        none_count += 1

            # 复制合并单元格
            for mc in src_ws.merged_cells.ranges:
                try:
                    new_ws.merge_cells(str(mc))
                except Exception:
                    pass

            # 复制列宽
            for col, dim in src_ws.column_dimensions.items():
                if dim.width:
                    new_ws.column_dimensions[col].width = dim.width

            # 若几乎全空（原文件无缓存值），回退到pandas读值
            if total_count == 0:
                df = pd.read_excel(str(src_path), sheet_name=sheet_name, header=None)
                df.to_excel(str(out_path), index=False, header=False)
            else:
                new_wb.save(str(out_path))
        else:
            # xls/csv 降级：用pandas读出再写xlsx
            if suffix == ".csv":
                df = pd.read_csv(str(src_path))
            else:
                df = pd.read_excel(str(src_path), sheet_name=sheet_name)
            df.to_excel(str(out_path), index=False, sheet_name=str(sheet_name)[:31])

        logger.info(f"单sheet导出: {source_file}[{sheet_name}] → {out_path.name}")
        return out_path

    # ---------- CSV ----------
    def _export_csv(self, df: pd.DataFrame, path: Path):
        df.to_csv(path, index=False, encoding="utf-8-sig")  # BOM头确保Windows Excel打开不乱码

    # ---------- Excel ----------
    def _export_excel(self, df: pd.DataFrame, path: Path):
        with pd.ExcelWriter(path, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="查询结果")
            self._auto_width(writer, df, "查询结果")

    def _auto_width(self, writer, df: pd.DataFrame, sheet_name: str):
        from openpyxl.utils import get_column_letter
        ws = writer.sheets[sheet_name]
        for i, col in enumerate(df.columns, 1):
            try:
                series_lens = df[col].astype(str).apply(lambda x: len(str(x)) if x is not None else 0)
                max_data_len = int(series_lens.max()) if len(series_lens) else 0
            except Exception:
                max_data_len = 10
            max_len = max(max_data_len, len(str(col)))
            # 中文字符宽度约为英文2倍，取保守估计
            width = min(max_len + 4, 50)
            ws.column_dimensions[get_column_letter(i)].width = width

    # ---------- PDF ----------
    def _export_pdf(self, df: pd.DataFrame, path: Path):
        """用 matplotlib 渲染表格型 PDF。
        说明：项目自带字体是 CFF/OTF 轮廓，fpdf2 子集化会乱码；
        matplotlib 字体管线对该字体支持正常，故 PDF 走 matplotlib 渲染。
        """
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib import font_manager
        from matplotlib.backends.backend_pdf import PdfPages

        # 注册中文字体（与图表导出同一管线，确保不乱码）
        if FONT_PATH.exists():
            try:
                font_manager.fontManager.addfont(str(FONT_PATH))
                plt.rcParams["font.family"] = font_manager.FontProperties(
                    fname=str(FONT_PATH)
                ).get_name()
            except Exception:
                logger.warning(f"中文字体注册失败: {FONT_PATH}")
        else:
            logger.warning(f"中文字体未找到: {FONT_PATH}，PDF可能出现乱码")
        plt.rcParams["axes.unicode_minus"] = False

        columns = [str(c) for c in df.columns]
        num_cols = len(columns)
        total = len(df)
        rows_per_page = 22  # A4横向一页约容纳的行数
        pages = max(1, (total + rows_per_page - 1) // rows_per_page)
        export_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        def _fmt(v) -> str:
            return ("" if pd.isna(v) else str(v))[:40]

        with PdfPages(str(path)) as pp:
            for pg in range(pages):
                sub = df.iloc[pg * rows_per_page:(pg + 1) * rows_per_page]
                fig, ax = plt.subplots(figsize=(11.69, 8.27))  # A4 横向(英寸)
                ax.axis("off")

                # 标题 + 导出时间
                fig.suptitle("财务数据查询结果", fontsize=15, fontweight="bold", y=0.97)
                fig.text(0.5, 0.92, f"导出时间: {export_time}",
                         ha="center", fontsize=8, color="gray")

                cell_text = [[_fmt(v) for v in row] for row in sub.values]
                if not cell_text:  # 空页保护
                    cell_text = [["" for _ in columns]]
                tbl = ax.table(cellText=cell_text, colLabels=columns,
                               loc="center", cellLoc="center")
                tbl.auto_set_font_size(False)
                tbl.set_fontsize(8)
                tbl.scale(1, 1.6)
                try:
                    tbl.auto_set_column_width(col=list(range(num_cols)))
                except Exception:
                    pass

                # 表头着色加粗
                for j in range(num_cols):
                    hc = tbl[0, j]
                    hc.set_facecolor("#DCE6F1")
                    hc.set_text_props(fontweight="bold")

                # 页脚统计
                fig.text(0.95, 0.04,
                         f"第 {pg + 1}/{pages} 页 · 共 {total} 行 × {num_cols} 列",
                         ha="right", fontsize=8, color="gray")

                pp.savefig(fig, bbox_inches="tight")
                plt.close(fig)

    # ---------- Word ----------
    def _export_docx(self, df: pd.DataFrame, path: Path):
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn

        doc = Document()

        # 设置默认中文字体
        style = doc.styles["Normal"]
        style.font.name = "Microsoft YaHei"  # Windows
        style.font.size = Pt(10)
        # 设置中文字体（兼容macOS显示"微软雅黑"不存在时回退到系统中文字体）
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

        # 标题
        title = doc.add_heading("财务数据查询结果", level=1)
        for run in title.runs:
            run.font.name = "Microsoft YaHei"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

        # 导出时间
        p = doc.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        p.style.font.size = Pt(9)
        p.style.font.color.rgb = RGBColor(128, 128, 128)

        # 表格
        columns = list(df.columns)
        table = doc.add_table(rows=1, cols=len(columns), style="Light Grid Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        header_row = table.rows[0]
        for i, col in enumerate(columns):
            cell = header_row.cells[i]
            cell.text = str(col)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.name = "Microsoft YaHei"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

        # 数据行
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, col in enumerate(columns):
                val = str(row[col]) if pd.notna(row[col]) else ""
                row_cells[i].text = val
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                        run.font.name = "Microsoft YaHei"
                        run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

        # 统计
        doc.add_paragraph("")
        summary = doc.add_paragraph(f"共 {len(df)} 行 × {len(columns)} 列")
        summary.style.font.size = Pt(9)

        doc.save(str(path))

    # ========== 文本问答导出（制度/RAG 类回答，无表格）==========
    def export_text(self, answer: str, question: str = "", sources: list = None,
                    snippets: list = None, fmt: str = "docx",
                    filename: str = "") -> Path:
        """导出文本类问答结果（问题 + 回答 + 引用来源 + 原文摘录）。"""
        if not answer or not str(answer).strip():
            raise ValueError("没有可导出的内容")
        if not filename:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"问答结果_{ts}"
        fmt = fmt.lower().strip(".")
        if fmt not in self.SUPPORTED_FORMATS:
            raise ValueError(f"不支持的导出格式: {fmt}")
        sources = sources or []
        snippets = snippets or []
        path = EXPORT_DIR / f"{filename}.{fmt}"

        if fmt == "docx":
            self._export_text_docx(answer, question, sources, snippets, path)
        elif fmt == "pdf":
            self._export_text_pdf(answer, question, sources, snippets, path)
        else:  # csv / xlsx
            self._export_text_table(answer, question, sources, fmt, path)

        logger.info(f"文本问答导出: {path.name}")
        return path

    @staticmethod
    def _fmt_source(s: dict) -> str:
        parts = [s.get("file", "")]
        if s.get("page"):
            parts.append(f"第{s['page']}页")
        if s.get("heading_path"):
            parts.append(s["heading_path"])
        elif s.get("section"):
            parts.append(s["section"])
        if s.get("relevance") is not None:
            parts.append(f"相关度{s['relevance']}")
        return " · ".join(str(p) for p in parts if p)

    @staticmethod
    def _wrap_cjk(text: str, width: int = 46) -> list:
        """按显示宽度换行（中文计2、英文计1），保留原有换行。"""
        out = []
        for para in str(text).split("\n"):
            if not para.strip():
                out.append("")
                continue
            cur, w = "", 0
            for ch in para:
                cw = 2 if ord(ch) > 0x2E80 else 1
                if w + cw > width:
                    out.append(cur)
                    cur, w = ch, cw
                else:
                    cur += ch
                    w += cw
            if cur:
                out.append(cur)
        return out

    def _export_text_table(self, answer, question, sources, fmt, path):
        rows = []
        if question:
            rows.append(("问题", str(question)))
        rows.append(("回答", str(answer)))
        for i, s in enumerate(sources, 1):
            rows.append((f"来源{i}", self._fmt_source(s)))
        df = pd.DataFrame(rows, columns=["字段", "内容"])
        if fmt == "csv":
            self._export_csv(df, path)
        else:
            self._export_excel(df, path)

    def _export_text_docx(self, answer, question, sources, snippets, path):
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn

        FONT = "Microsoft YaHei"

        def _cn(run):
            run.font.name = FONT
            run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = FONT
        style.font.size = Pt(10.5)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

        h = doc.add_heading("财务知识库问答结果", level=1)
        for r in h.runs:
            _cn(r)
        p = doc.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        for r in p.runs:
            _cn(r)
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(128, 128, 128)

        if question:
            hq = doc.add_heading("问题", level=2)
            for r in hq.runs:
                _cn(r)
            pq = doc.add_paragraph(str(question))
            for r in pq.runs:
                _cn(r)

        ha = doc.add_heading("回答", level=2)
        for r in ha.runs:
            _cn(r)
        for line in str(answer).split("\n"):
            pa = doc.add_paragraph(line)
            for r in pa.runs:
                _cn(r)

        if sources:
            hs = doc.add_heading("引用来源", level=2)
            for r in hs.runs:
                _cn(r)
            for s in sources:
                ps = doc.add_paragraph(self._fmt_source(s), style="List Bullet")
                for r in ps.runs:
                    _cn(r)

        if snippets:
            hsn = doc.add_heading("原文摘录", level=2)
            for r in hsn.runs:
                _cn(r)
            for i, sn in enumerate(snippets, 1):
                tag = f"摘录{i}（{sn.get('file','')}"
                if sn.get("page"):
                    tag += f", 第{sn['page']}页"
                tag += "）"
                ph = doc.add_paragraph(tag)
                for r in ph.runs:
                    r.font.bold = True
                    _cn(r)
                pt = doc.add_paragraph(sn.get("text", ""))
                for r in pt.runs:
                    _cn(r)

        doc.save(str(path))

    def _export_text_pdf(self, answer, question, sources, snippets, path):
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib import font_manager
        from matplotlib.backends.backend_pdf import PdfPages

        if FONT_PATH.exists():
            try:
                font_manager.fontManager.addfont(str(FONT_PATH))
                plt.rcParams["font.family"] = font_manager.FontProperties(
                    fname=str(FONT_PATH)
                ).get_name()
            except Exception:
                logger.warning(f"中文字体注册失败: {FONT_PATH}")
        plt.rcParams["axes.unicode_minus"] = False

        lines = []
        if question:
            lines.append("【问题】")
            lines += self._wrap_cjk(question, 46)
            lines.append("")
        lines.append("【回答】")
        lines += self._wrap_cjk(answer, 46)
        if sources:
            lines += ["", "【引用来源】"]
            for s in sources:
                lines += self._wrap_cjk("· " + self._fmt_source(s), 46)
        if snippets:
            lines += ["", "【原文摘录】"]
            for i, sn in enumerate(snippets, 1):
                tag = f"摘录{i}（{sn.get('file','')}"
                if sn.get("page"):
                    tag += f", 第{sn['page']}页"
                tag += "）："
                lines += self._wrap_cjk(tag + sn.get("text", ""), 46)

        lines_per_page = 40
        pages = max(1, (len(lines) + lines_per_page - 1) // lines_per_page)
        with PdfPages(str(path)) as pp:
            for pg in range(pages):
                chunk = lines[pg * lines_per_page:(pg + 1) * lines_per_page]
                fig = plt.figure(figsize=(8.27, 11.69))  # A4 纵向
                fig.text(0.5, 0.965, "财务知识库问答结果", ha="center",
                         fontsize=14, fontweight="bold")
                y = 0.93
                for ln in chunk:
                    fig.text(0.08, y, ln, ha="left", va="top", fontsize=10.5)
                    y -= 0.0225
                fig.text(0.92, 0.025, f"第 {pg + 1}/{pages} 页",
                         ha="right", fontsize=8, color="gray")
                pp.savefig(fig)
                plt.close(fig)
