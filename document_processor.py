"""文档接入层：三通道PDF解析 + 结构还原 + 元数据保留"""

import re
import statistics
from pathlib import Path
from dataclasses import dataclass, field
from datetime import datetime

import fitz
import pandas as pd
from docx import Document as DocxDocument
from loguru import logger


# ============================================================
# 数据结构
# ============================================================
@dataclass
class TextBlock:
    content: str
    source_file: str
    page_num: int = 0
    section: str = ""          # 所属章节标题
    heading_path: str = ""     # 章节路径，如 "第一章 > 第一节 > 1.1"
    block_type: str = "body"   # heading / body / footnote / toc
    doc_type: str = "text"     # text / regulation / policy
    font_info: str = ""        # 字体/字号摘要


@dataclass
class TableBlock:
    dataframe: pd.DataFrame
    source_file: str
    sheet_name: str = ""
    page_num: int = 0
    table_index: int = 0
    caption: str = ""          # 表格标题/上方文字
    metadata: dict = field(default_factory=dict)


@dataclass
class ImageBlock:
    """PDF中嵌入的图片/截图"""
    image_path: str            # 本地保存路径
    source_file: str
    page_num: int = 0
    image_index: int = 0
    width: int = 0
    height: int = 0
    caption: str = ""          # 图片上下方的文字（可选）


@dataclass
class ParseResult:
    texts: list[TextBlock] = field(default_factory=list)
    tables: list[TableBlock] = field(default_factory=list)
    images: list[ImageBlock] = field(default_factory=list)
    file_name: str = ""
    file_type: str = ""
    parse_time: str = ""
    pdf_type: str = ""         # native / scanned / mixed


# ============================================================
# PDF 结构解析器
# ============================================================
class PDFStructureParser:
    """从 PyMuPDF dict 输出中还原文档结构"""

    # 页眉页脚占页面高度的比例阈值
    HEADER_RATIO = 0.08
    FOOTER_RATIO = 0.92
    # 脚注字体相对正文缩小比例
    FOOTNOTE_SIZE_RATIO = 0.75
    # 目录页关键词
    TOC_KEYWORDS = ["目录", "目 录", "CONTENTS", "TABLE OF CONTENTS"]

    def __init__(self):
        self._body_font_size = None  # 正文基准字号

    def detect_pdf_type(self, doc: fitz.Document) -> str:
        """检测PDF类型：native / scanned / mixed"""
        total_pages = len(doc)
        text_pages = 0
        image_pages = 0

        sample = range(0, total_pages, max(1, total_pages // 10))  # 抽样检测
        for i in sample:
            page = doc[i]
            text = page.get_text("text").strip()
            images = page.get_images()
            if len(text) > 50:
                text_pages += 1
            if images and len(text) < 20:
                image_pages += 1

        if image_pages == 0:
            return "native"
        elif text_pages == 0:
            return "scanned"
        else:
            return "mixed"

    def get_body_font_size(self, doc: fitz.Document) -> float:
        """统计正文基准字号（出现频率最高的字号）"""
        size_count = {}
        sample_pages = min(10, len(doc))
        for i in range(sample_pages):
            page = doc[i]
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block["type"] != 0:  # 非文字块
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        size = round(span["size"], 1)
                        text = span["text"].strip()
                        if len(text) > 5:  # 忽略太短的片段
                            size_count[size] = size_count.get(size, 0) + len(text)

        if not size_count:
            return 12.0
        self._body_font_size = max(size_count, key=size_count.get)
        return self._body_font_size

    def is_heading(self, span: dict) -> bool:
        """判断是否是标题"""
        if self._body_font_size is None:
            return False
        size = span["size"]
        text = span["text"].strip()
        flags = span.get("flags", 0)
        is_bold = bool(flags & 2 ** 4)  # bit 4 = bold

        # 字号明显大于正文
        if size > self._body_font_size * 1.15:
            return True
        # 加粗且字号不小于正文
        if is_bold and size >= self._body_font_size and len(text) < 80:
            return True
        # 章节编号模式
        if re.match(r"^(第[一二三四五六七八九十百\d]+[章节条款篇]|[\d]+[\.\s])", text):
            return True
        return False

    def is_footnote(self, block: dict, page_height: float) -> bool:
        """判断是否是脚注（位于页面底部 + 字号偏小）"""
        if block["type"] != 0:
            return False
        bbox = block["bbox"]
        if bbox[1] < page_height * self.FOOTER_RATIO:
            return False
        # 检查字号
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                if self._body_font_size and span["size"] < self._body_font_size * self.FOOTNOTE_SIZE_RATIO:
                    return True
        return False

    def is_header_footer(self, block: dict, page_height: float) -> bool:
        """判断是否是页眉页脚"""
        bbox = block["bbox"]
        # 页眉区域
        if bbox[3] < page_height * self.HEADER_RATIO:
            return True
        # 页脚区域（纯页码）
        if bbox[1] > page_height * self.FOOTER_RATIO:
            text = ""
            if block["type"] == 0:
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text += span["text"]
            text = text.strip()
            # 纯数字或很短的页脚
            if re.match(r"^[\d\s\-—/页第共of]+$", text) or len(text) < 10:
                return True
        return False

    def is_garbled(self, text: str) -> bool:
        """检测一段文本是否含有PDF截图CID乱码"""
        if not text or len(text) < 20:
            return False

        # 整段全乱码检测
        if self._is_chunk_garbled(text):
            return True

        # 段内夹杂检测：长文本中可疑token占比过高
        if len(text) > 60:
            tokens = text.split()
            tokens = [t for t in tokens if len(t) >= 1]
            if len(tokens) >= 10:
                suspicious = sum(1 for t in tokens if self._is_suspicious_token(t))
                # 可疑token占比 >= 25%  或  绝对数量 >= 8 → 视为含乱码
                if suspicious / len(tokens) >= 0.25 or suspicious >= 8:
                    return True
        return False

    def clean_sentences(self, text: str) -> str:
        """按句子粒度过滤：把段切成多个语义单元，逐个判断，只保留正常句子。
        语义单元定义：以中英文句号/感叹号/问号/换行/分号结尾。
        """
        if not text or len(text) < 20:
            return text

        # 按句末标点+换行 切分成句子（保留分隔符）
        # 使用前瞻断言保留标点
        sentences = re.split(r"(?<=[。！？；\n])", text)
        sentences = [s for s in sentences if s.strip()]

        if len(sentences) < 2:
            # 没切出多个句子，回退尾部截断
            return self.clean_garbled_tail(text)

        # 逐句判断
        kept = []
        for s in sentences:
            s_clean = s.strip()
            if len(s_clean) < 5:
                # 太短的片段跟随上下文：如果上一句已保留就一起留
                if kept:
                    kept.append(s)
                continue
            # 单句乱码判断（用更宽容的阈值，单句信息少）
            if self._is_sentence_garbled(s_clean):
                continue
            kept.append(s)

        if not kept:
            return ""
        result = "".join(kept).strip()
        # 如果保留下来的内容不到原文40%，认为破坏严重不动
        if len(result) < len(text) * 0.4 and len(result) < 30:
            return text
        return result

    def _is_line_garbled(self, line: str) -> bool:
        """单行乱码判断（PyMuPDF返回的物理行级别）"""
        clean = re.sub(r"\s+", "", line)
        n = len(clean)
        if n < 2:
            return False

        cjk = sum(1 for c in clean if "一" <= c <= "鿿")
        ascii_letters = sum(1 for c in clean if c.isascii() and c.isalpha())
        digits = sum(1 for c in clean if c.isdigit())
        punct = sum(1 for c in clean if c in "，。；：、！？,.;:!?（）()【】[]《》<>—-…")
        meaningful = cjk + ascii_letters + digits + punct
        others = n - meaningful

        # 罕见符号比例高
        if others / n > 0.15 and others >= 2:
            return True

        # 全是数字+少量字母（如 "1025108731" 单独一行）
        if cjk == 0 and ascii_letters <= 2 and digits >= 6:
            return True

        # 单字 + 字母数字混合的可疑组合（如 "税种 单5为立"）
        tokens = line.split()
        if len(tokens) >= 2:
            susp = sum(1 for t in tokens if self._is_suspicious_token(t))
            # 短行（<60字符）中，可疑token>=2 或占比>=40%
            if n <= 60:
                if susp >= 2 or (len(tokens) > 0 and susp / len(tokens) >= 0.40):
                    return True
            else:
                # 长行：可疑占比>=35%
                if susp / len(tokens) >= 0.35:
                    return True

        # 中文占比极低且 ASCII 字母大小写混合短词为主
        if 8 <= n <= 80 and cjk / n < 0.20:
            mixed_tokens = [t for t in tokens if t.isascii() and t.isalpha()
                            and any(c.isupper() for c in t) and any(c.islower() for c in t)]
            if len(mixed_tokens) >= 2:
                return True

        # 大量"中文+字母数字"紧密混合的token（如 "立5业", "下l", "百m"）
        cjk_mix = sum(1 for t in tokens
                      if 2 <= len(t) <= 6
                      and re.search(r"[A-Za-z0-9]", t)
                      and re.search(r"[一-鿿]", t))
        if cjk_mix >= 3:
            return True

        # 大量"全大写ASCII短串"行（典型CID字形）
        # 如 "IS P FNSUUUODI TTI2IITUI FSTNUUIT" "FSUIUIII P"
        all_caps_short = sum(1 for t in tokens
                              if t.isascii() and t.isalpha()
                              and 1 <= len(t) <= 10
                              and t.isupper())
        if all_caps_short >= 3 and all_caps_short / max(len(tokens), 1) >= 0.5:
            return True

        # 超短行（1-3 token）特殊规则
        if 1 <= len(tokens) <= 3:
            # 含可疑token直接丢（注意：可疑token不包括正常的事务码/序号）
            susp = sum(1 for t in tokens if self._is_suspicious_token(t))
            if susp >= 1:
                return True
            # 全是全大写无含义短串（且非业务缩写白名单）
            whitelist = {"SAP", "ERP", "OA", "HR", "CEO", "CFO", "VP", "FY",
                         "Q1", "Q2", "Q3", "Q4", "AR", "AP", "GL",
                         "USD", "CNY", "EUR", "USA", "HK", "OK", "NO", "ID"}
            alpha_tokens = [t for t in tokens if t.isascii() and t.isalpha()]
            if alpha_tokens:
                upper_alpha = [t for t in alpha_tokens if t.isupper() and 2 <= len(t) <= 10]
                if upper_alpha == alpha_tokens and not all(t in whitelist for t in alpha_tokens):
                    # 排除"OA系统"这种带中文后缀的情况：检查整行
                    if cjk == 0:
                        return True

        return False

    def _is_sentence_garbled(self, s: str) -> bool:
        """单句乱码判断（更宽容）"""
        clean = re.sub(r"\s+", "", s)
        n = len(clean)
        if n < 5:
            return False

        cjk = sum(1 for c in clean if "一" <= c <= "鿿")
        ascii_letters = sum(1 for c in clean if c.isascii() and c.isalpha())
        digits = sum(1 for c in clean if c.isdigit())
        punct = sum(1 for c in clean if c in "，。；：、！？,.;:!?（）()【】[]《》<>—-…")
        meaningful = cjk + ascii_letters + digits + punct
        others = n - meaningful

        # 罕见符号占比 >15%
        if n >= 10 and others / n > 0.15:
            return True

        # 短句中文 < 20% 且字母多 → 大概率乱码
        if 10 <= n <= 60 and cjk / n < 0.20 and ascii_letters > 5:
            tokens = s.split()
            if tokens:
                susp = sum(1 for t in tokens if self._is_suspicious_token(t))
                if susp >= 2 or (len(tokens) > 0 and susp / len(tokens) > 0.3):
                    return True

        # 可疑token绝对数检测
        tokens = s.split()
        if len(tokens) >= 3:
            susp = sum(1 for t in tokens if self._is_suspicious_token(t))
            if susp >= 3 and susp / len(tokens) >= 0.4:
                return True

        return False

    def clean_garbled_tail(self, text: str) -> str:
        """从段末截断乱码尾巴；保留前半正常内容。
        策略：找文本中所有"完整句结束位置"（中文句号/感叹号/问号），
        对每个候选切点，检查切点之后的尾部是否乱码，是则截断到该切点。
        """
        if not text or len(text) < 30:
            return text

        # 找所有句末标点位置（中文优先）
        cut_candidates = []
        for m in re.finditer(r"[。！？；\.\!\?]", text):
            cut_candidates.append(m.end())

        if not cut_candidates:
            return text

        # 从后往前找：找最后一个"切点之后是乱码"的位置
        for cut_pos in reversed(cut_candidates):
            tail = text[cut_pos:].strip()
            head = text[:cut_pos].strip()

            # 尾部过短（<10字符）跳过，可能只是句末小尾巴
            if len(tail) < 10:
                continue
            # 头部过短（<20字符）跳过，避免截掉过多
            if len(head) < 20:
                continue

            # 判断尾部是否含可疑token序列
            tail_tokens = tail.split()
            if len(tail_tokens) < 3:
                continue
            tail_suspicious = sum(1 for t in tail_tokens if self._is_suspicious_token(t))
            # 尾部可疑token >=2 且占比 >=25%，就截断
            if tail_suspicious >= 2 and tail_suspicious / len(tail_tokens) >= 0.25:
                return head + "…"

        return text

    def _is_suspicious_token(self, tok: str) -> bool:
        """单个token是否疑似CID乱码字形"""
        if not tok or len(tok) < 2:
            return False

        has_alnum = bool(re.search(r"[A-Za-z0-9]", tok))
        has_cjk = bool(re.search(r"[一-鿿]", tok))
        # 注意：故意不把 . / - 视为 noise 符号，避免误杀 "1.1章节"、"a-b" 这类正常token
        has_symbol = bool(re.search(r"[~`!@#$%^&*+=\\|/?<>\[\]\{\}『』〔〕《》]", tok))

        # 1) ASCII大小写混合短词（如 GtetSapWT, REappr, CTprc）
        if tok.isascii() and 2 <= len(tok) <= 10:
            letters = [c for c in tok if c.isalpha()]
            if len(letters) >= 2:
                has_upper = any(c.isupper() for c in letters)
                has_lower = any(c.islower() for c in letters)
                if has_upper and has_lower:
                    inner_upper = sum(1 for i, c in enumerate(tok) if i > 0 and c.isupper())
                    if inner_upper >= 1:
                        return True

        # 2) 短token：汉字+字母数字混合（如 "正TSIpRUUo", "下l", "百m"）
        # 注意：cjk_count == 1 才算可疑，避免误杀 "OA系统"(cjk=2)、"PVH台"(cjk=2) 等正常组合
        if 2 <= len(tok) <= 6 and has_alnum and has_cjk:
            cjk_count = sum(1 for c in tok if "一" <= c <= "鿿")
            if cjk_count == 1:
                return True

        # 3) 罕见符号
        if re.search(r"[『』〔〕《》\|\[\]\{\}]", tok):
            return True

        # 4) 全大写短串+数字混合（如 TTO10179）
        if tok.isascii() and 4 <= len(tok) <= 10:
            digits = sum(1 for c in tok if c.isdigit())
            uppers = sum(1 for c in tok if c.isupper())
            lowers = sum(1 for c in tok if c.islower())
            if digits >= 1 and uppers >= 2 and uppers > lowers:
                return True

        # 5) 长token：含【汉字+字母+符号】三种紧密混合（如 "亩-鬟舌咀量溟t_-痄止醒p"）
        if len(tok) >= 5 and has_cjk and has_alnum and has_symbol:
            return True

        # 6) 短token：单个中文+多个字母数字（如 "正TSIpRUUo"）
        # 严格条件：cjk_count == 1，避免误杀 "OA系统"、"1.1应付账款" 等正常组合
        if 4 <= len(tok) <= 10 and has_cjk and has_alnum:
            cjk_count = sum(1 for c in tok if "一" <= c <= "鿿")
            alnum_count = sum(1 for c in tok if c.isascii() and (c.isalnum()))
            if cjk_count == 1 and alnum_count >= 3:
                if re.search(r"[一-鿿][A-Za-z]|[A-Za-z][一-鿿]", tok):
                    return True

        # 7) 短全大写或短大小写混合，含罕见标点（如 "T[5"）
        if len(tok) <= 6 and has_symbol and re.search(r"[A-Za-z]", tok):
            return True

        return False

    def _is_chunk_garbled(self, text: str) -> bool:
        clean = re.sub(r"\s+", "", text)
        n = len(clean)
        if n < 20:
            return False

        cjk = sum(1 for c in clean if "一" <= c <= "鿿")
        ascii_letters = sum(1 for c in clean if c.isascii() and c.isalpha())
        digits = sum(1 for c in clean if c.isdigit())
        punct = sum(1 for c in clean if c in "，。；：、！？,.;:!?（）()【】[]《》<>—-…")
        meaningful = cjk + ascii_letters + digits + punct
        others = n - meaningful

        if cjk / n < 0.15 and ascii_letters / n > 0.4:
            tokens = re.findall(r"[A-Za-z]+", clean)
            if tokens:
                avg_len = sum(len(t) for t in tokens) / len(tokens)
                mixed_case = sum(1 for t in tokens if any(c.isupper() for c in t) and any(c.islower() for c in t))
                if avg_len < 5 and mixed_case / len(tokens) > 0.3:
                    return True

        if others / n > 0.20:
            return True

        alternating = len(re.findall(r"[A-Za-z][一-鿿][A-Za-z]", clean))
        if alternating > 5 and alternating / n > 0.05:
            return True

        return False

    def is_toc_page(self, page_text: str) -> bool:
        """判断是否是目录页"""
        first_200 = page_text[:200].strip()
        for kw in self.TOC_KEYWORDS:
            if kw in first_200:
                return True
        # 大量 "...数字" 模式
        dot_pattern = re.findall(r"[\.·…]{3,}\s*\d+", page_text)
        if len(dot_pattern) > 5:
            return True
        return False

    def extract_structured_blocks(self, page: fitz.Page, page_num: int,
                                   source_file: str) -> tuple[list[TextBlock], str]:
        """从单页提取结构化文本块，返回 (blocks, 最后一个标题)"""
        page_dict = page.get_text("dict")
        page_height = page.rect.height
        blocks = []

        for block in page_dict["blocks"]:
            if block["type"] != 0:  # 跳过图片块
                continue

            if self.is_header_footer(block, page_height):
                continue

            # 按行级别提取并过滤乱码行
            block_lines = []
            block_spans = []
            for line in block.get("lines", []):
                line_text = ""
                line_spans = []
                for span in line.get("spans", []):
                    line_text += span["text"]
                    line_spans.append(span)

                line_text = line_text.strip()
                if not line_text:
                    continue

                # 行级乱码过滤
                if self._is_line_garbled(line_text):
                    continue

                block_lines.append(line_text)
                block_spans.extend(line_spans)

            block_text = "\n".join(block_lines).strip()
            if not block_text or len(block_text) < 5:
                continue

            # 再做一次句子级兜底过滤
            cleaned = self.clean_sentences(block_text)
            if not cleaned or len(cleaned) < 10:
                logger.debug(f"丢弃乱码块: {block_text[:60]}...")
                continue
            block_text = cleaned

            # 判断块类型
            if self.is_footnote(block, page_height):
                block_type = "footnote"
            elif block_spans and self.is_heading(block_spans[0]):
                block_type = "heading"
            else:
                block_type = "body"

            # 字体摘要
            font_sizes = [s["size"] for s in block_spans if s["text"].strip()]
            font_info = f"size={round(statistics.mean(font_sizes), 1)}" if font_sizes else ""

            blocks.append(TextBlock(
                content=block_text,
                source_file=source_file,
                page_num=page_num,
                block_type=block_type,
                font_info=font_info,
            ))

        return blocks


# ============================================================
# OCR 处理器
# ============================================================
class OCRProcessor:
    """扫描件 OCR 提取文字"""

    def __init__(self):
        self._reader = None

    def _get_reader(self):
        if self._reader is None:
            import easyocr
            self._reader = easyocr.Reader(["ch_sim", "en"], gpu=False)
        return self._reader

    def ocr_page(self, page: fitz.Page, source_file: str, page_num: int) -> list[TextBlock]:
        """对单页进行 OCR"""
        pix = page.get_pixmap(dpi=200)
        img_data = pix.tobytes("png")

        import io
        from PIL import Image
        img = Image.open(io.BytesIO(img_data))
        import numpy as np
        img_array = np.array(img)

        reader = self._get_reader()
        # paragraph=True 返回 [(bbox, text)] 二元组；paragraph=False 返回 [(bbox, text, conf)] 三元组
        results = reader.readtext(img_array, paragraph=True)

        blocks = []
        for item in results:
            # 兼容 2/3 元组
            if len(item) == 3:
                bbox, text, conf = item
            elif len(item) == 2:
                bbox, text = item
                conf = None
            else:
                continue

            text = (text or "").strip()
            if len(text) < 3:
                continue

            font_info = f"ocr_conf={conf:.2f}" if conf is not None else "ocr"
            blocks.append(TextBlock(
                content=text,
                source_file=source_file,
                page_num=page_num,
                block_type="ocr",
                font_info=font_info,
            ))

        return blocks


# ============================================================
# 主处理器
# ============================================================
class DocumentProcessor:
    """统一文档解析入口：三通道PDF + Word/Excel/CSV"""

    SUPPORTED = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".xlsm"}

    def __init__(self):
        self.pdf_parser = PDFStructureParser()
        self._ocr = None  # 延迟初始化

    @property
    def ocr(self):
        if self._ocr is None:
            self._ocr = OCRProcessor()
        return self._ocr

    def parse(self, file_path: str | Path) -> ParseResult:
        path = Path(file_path)
        if path.suffix.lower() not in self.SUPPORTED:
            raise ValueError(f"不支持的文件格式: {path.suffix}")

        result = ParseResult(
            file_name=path.name,
            file_type=path.suffix.lower(),
            parse_time=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        )

        handler = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".doc": self._parse_docx,
            ".xlsx": self._parse_excel,
            ".xls": self._parse_excel,
            ".xlsm": self._parse_excel,
            ".csv": self._parse_csv,
        }

        handler[path.suffix.lower()](path, result)

        # 构建章节路径
        self._build_heading_paths(result.texts)

        logger.info(
            f"解析完成: {path.name} [类型:{result.pdf_type or result.file_type}] "
            f"→ 文本块{len(result.texts)}个, 表格{len(result.tables)}个, 图片{len(result.images)}张"
        )
        return result

    # ---------- PDF 三通道 ----------
    def _parse_pdf(self, path: Path, result: ParseResult):
        from config import DATA_DIR
        img_dir = DATA_DIR / "extracted_images"
        img_dir.mkdir(parents=True, exist_ok=True)

        doc = fitz.open(str(path))
        pdf_type = self.pdf_parser.detect_pdf_type(doc)
        result.pdf_type = pdf_type
        logger.info(f"PDF类型检测: {path.name} → {pdf_type}")

        body_size = self.pdf_parser.get_body_font_size(doc)
        logger.info(f"正文基准字号: {body_size}")

        # 文件名安全化：用作图片文件前缀
        import re as _re
        safe_stem = _re.sub(r"[^\w一-鿿\.\-]", "_", path.stem)

        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text("text").strip()

            # 跳过目录页
            if self.pdf_parser.is_toc_page(page_text):
                logger.debug(f"跳过目录页: 第{page_num}页")
                continue

            # 提取嵌入图片（截图、插图）
            for img_idx, img_info in enumerate(page.get_images(full=True)):
                try:
                    xref = img_info[0]
                    base_img = doc.extract_image(xref)
                    img_bytes = base_img["image"]
                    img_ext = base_img.get("ext", "png")
                    img_w = base_img.get("width", 0)
                    img_h = base_img.get("height", 0)

                    # 过滤极小图片（如装饰图标），宽高都<60视为非有效截图
                    if img_w < 60 and img_h < 60:
                        continue

                    img_filename = f"{safe_stem}_p{page_num:03d}_img{img_idx:02d}.{img_ext}"
                    img_path = img_dir / img_filename
                    img_path.write_bytes(img_bytes)

                    result.images.append(ImageBlock(
                        image_path=str(img_path),
                        source_file=path.name,
                        page_num=page_num,
                        image_index=img_idx,
                        width=img_w,
                        height=img_h,
                    ))
                except Exception as e:
                    logger.warning(f"图片提取失败 p{page_num}_img{img_idx}: {e}")

            # 提取表格（所有类型都尝试）
            tables = page.find_tables()
            for idx, table in enumerate(tables):
                df = table.to_pandas()
                if df is not None and not df.empty and len(df) > 1:
                    df = self._clean_table_headers(df)
                    # 尝试提取表格上方的标题
                    caption = self._extract_table_caption(page, table.bbox)
                    result.tables.append(TableBlock(
                        dataframe=df,
                        source_file=path.name,
                        page_num=page_num,
                        table_index=idx,
                        caption=caption,
                    ))

            # 根据PDF类型选择文字提取方式
            if pdf_type == "native" or (pdf_type == "mixed" and len(page_text) > 50):
                # 原生文字：结构化提取
                blocks = self.pdf_parser.extract_structured_blocks(
                    page, page_num, path.name
                )
                result.texts.extend(blocks)

            elif pdf_type == "scanned" or (pdf_type == "mixed" and len(page_text) <= 50):
                # 扫描件：OCR
                ocr_blocks = self.ocr.ocr_page(page, path.name, page_num)
                result.texts.extend(ocr_blocks)

        doc.close()

    def _extract_table_caption(self, page: fitz.Page, table_bbox: tuple) -> str:
        """提取表格正上方的文字作为表格标题"""
        page_dict = page.get_text("dict")
        table_top = table_bbox[1]
        candidates = []

        for block in page_dict["blocks"]:
            if block["type"] != 0:
                continue
            block_bottom = block["bbox"][3]
            # 在表格上方 50pt 以内
            if table_top - 50 < block_bottom < table_top:
                text = ""
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text += span["text"]
                text = text.strip()
                if text and len(text) < 100:
                    candidates.append((block_bottom, text))

        if candidates:
            candidates.sort(key=lambda x: x[0], reverse=True)
            return candidates[0][1]
        return ""

    # ---------- 章节路径构建 ----------
    def _build_heading_paths(self, texts: list[TextBlock]):
        """为所有文本块填充 section 和 heading_path"""
        heading_stack = []  # [(level, title)]
        current_section = ""
        current_path = ""

        for block in texts:
            if block.block_type == "heading":
                level = self._guess_heading_level(block)
                title = block.content.strip().split("\n")[0][:80]

                # 弹出同级或更低级的标题
                while heading_stack and heading_stack[-1][0] >= level:
                    heading_stack.pop()
                heading_stack.append((level, title))

                current_section = title
                current_path = " > ".join(h[1] for h in heading_stack)

            block.section = current_section
            block.heading_path = current_path

    def _guess_heading_level(self, block: TextBlock) -> int:
        """根据字号和内容猜测标题层级"""
        text = block.content.strip()
        # 第X章 = 1级
        if re.match(r"^第[一二三四五六七八九十百\d]+章", text):
            return 1
        # 第X节 = 2级
        if re.match(r"^第[一二三四五六七八九十百\d]+节", text):
            return 2
        # X.X 编号 = 按点数判断
        m = re.match(r"^(\d+(?:\.\d+)*)", text)
        if m:
            return m.group(1).count(".") + 1
        # 按字号
        if "size=" in block.font_info:
            size = float(block.font_info.split("size=")[1].split(",")[0])
            if size > 18:
                return 1
            elif size > 14:
                return 2
            else:
                return 3
        return 3

    # ---------- Word ----------
    def _parse_docx(self, path: Path, result: ParseResult):
        doc = DocxDocument(str(path))
        current_section = ""
        heading_stack = []
        text_buffer = []
        # docx 无页码概念，统一用第1页，便于图片做文件级关联
        DOCX_PAGE = 1

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""

            if style_name.startswith("Heading"):
                # 先保存之前累积的文本
                if text_buffer:
                    result.texts.append(TextBlock(
                        content="\n".join(text_buffer),
                        source_file=path.name,
                        page_num=DOCX_PAGE,
                        section=current_section,
                        heading_path=" > ".join(h[1] for h in heading_stack),
                        block_type="body",
                    ))
                    text_buffer = []

                # 解析标题层级
                level = int(re.search(r"\d+", style_name).group()) if re.search(r"\d+", style_name) else 1
                title = para.text.strip()
                while heading_stack and heading_stack[-1][0] >= level:
                    heading_stack.pop()
                heading_stack.append((level, title))
                current_section = title

                result.texts.append(TextBlock(
                    content=title,
                    source_file=path.name,
                    page_num=DOCX_PAGE,
                    section=current_section,
                    heading_path=" > ".join(h[1] for h in heading_stack),
                    block_type="heading",
                ))
            elif para.text.strip():
                text_buffer.append(para.text.strip())

        if text_buffer:
            result.texts.append(TextBlock(
                content="\n".join(text_buffer),
                source_file=path.name,
                page_num=DOCX_PAGE,
                section=current_section,
                heading_path=" > ".join(h[1] for h in heading_stack),
                block_type="body",
            ))

        # Word 表格
        for idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if len(rows) > 1:
                df = pd.DataFrame(rows[1:], columns=rows[0])
                result.tables.append(TableBlock(
                    dataframe=df,
                    source_file=path.name,
                    page_num=DOCX_PAGE,
                    table_index=idx,
                ))

        # 提取 docx 嵌入图片（截图、插图）
        self._extract_docx_images(doc, path, result, DOCX_PAGE)

    def _extract_docx_images(self, doc, path: Path, result: ParseResult, page_num: int):
        """从 docx 的关系部件中提取所有嵌入图片"""
        from config import DATA_DIR
        import re as _re
        img_dir = DATA_DIR / "extracted_images"
        img_dir.mkdir(parents=True, exist_ok=True)
        safe_stem = _re.sub(r"[^\w一-鿿\.\-]", "_", path.stem)

        img_idx = 0
        try:
            rels = doc.part.rels
        except Exception as e:
            logger.warning(f"docx 关系部件读取失败: {e}")
            return

        for rel in rels.values():
            if "image" not in rel.reltype:
                continue
            try:
                blob = rel.target_part.blob
                content_type = getattr(rel.target_part, "content_type", "")
                ext = content_type.split("/")[-1] if "/" in content_type else "png"
                if ext in ("jpeg", "jpg"):
                    ext = "jpg"
                elif ext not in ("png", "gif", "bmp", "tiff", "emf", "wmf"):
                    ext = "png"

                # 过滤极小装饰图（用 Pillow 读尺寸，失败则按字节大小过滤）
                w = h = 0
                try:
                    from PIL import Image
                    import io as _io
                    with Image.open(_io.BytesIO(blob)) as im:
                        w, h = im.size
                    if w < 60 and h < 60:
                        continue
                except Exception:
                    if len(blob) < 3000:  # <3KB 多为图标
                        continue

                img_filename = f"{safe_stem}_docx_img{img_idx:03d}.{ext}"
                img_path = img_dir / img_filename
                img_path.write_bytes(blob)

                result.images.append(ImageBlock(
                    image_path=str(img_path),
                    source_file=path.name,
                    page_num=page_num,
                    image_index=img_idx,
                    width=w,
                    height=h,
                ))
                img_idx += 1
            except Exception as e:
                logger.warning(f"docx 图片提取失败 img{img_idx}: {e}")

        if img_idx:
            logger.info(f"docx 嵌入图片提取: {img_idx}张 ← {path.name}")

    # ---------- Excel / CSV ----------
    # 财务报表常见列头关键词（用于智能识别真正的列头行）
    HEADER_KEYWORDS = [
        "期末数", "期初数", "年初数", "年末数", "本期数", "上期数", "本年", "上年",
        "本期发生额", "上期发生额", "本年累计", "上年累计", "上年同期",
        "金额", "数量", "单价", "合计", "余额", "借方", "贷方",
        "科目编码", "科目名称", "科目代码", "科目", "项目",
        "注释号", "附注", "行次", "编号", "凭证号",
        "资产", "负债", "所有者权益", "收入", "费用", "成本",
        "部门", "员工", "客户", "供应商", "产品", "区域",
    ]
    # 描述性行关键词（用于跳过）
    DESC_KEYWORDS = [
        "编制单位", "单位：", "单位:", "会企", "金额单位",
        "公司名称", "报表日期", "编制日期", "审核",
    ]

    def _parse_excel(self, path: Path, result: ParseResult):
        xls = pd.ExcelFile(str(path))
        for sheet_name in xls.sheet_names:
            # 不指定header，先全部读进来，再智能定位
            raw_df = pd.read_excel(xls, sheet_name=sheet_name, header=None)
            raw_df = raw_df.dropna(how="all").reset_index(drop=True)
            if raw_df.empty:
                continue

            header_row = self._detect_header_row(raw_df)
            df = self._extract_table_from_header(raw_df, header_row)

            if df.empty or len(df.columns) == 0:
                continue

            result.tables.append(TableBlock(
                dataframe=df,
                source_file=path.name,
                sheet_name=sheet_name,
                metadata={
                    "total_rows": len(df),
                    "columns": list(df.columns),
                    "header_row": header_row,
                },
            ))

    def _detect_header_row(self, raw_df: pd.DataFrame, max_scan: int = 15) -> int:
        """智能识别真正的列头所在行，扫描前 max_scan 行"""
        scan_end = min(max_scan, len(raw_df))
        best_row = 0
        best_score = -1

        for i in range(scan_end):
            row = raw_df.iloc[i]
            row_str = " ".join(str(v) for v in row if pd.notna(v))

            # 跳过描述行
            if any(kw in row_str for kw in self.DESC_KEYWORDS):
                continue

            # 跳过纯数字行（数据行）
            non_null = [v for v in row if pd.notna(v) and str(v).strip()]
            if not non_null:
                continue
            num_count = sum(1 for v in non_null if self._is_numeric(v))
            if num_count > len(non_null) * 0.5:
                continue

            # 评分：含表头关键词越多越像表头
            kw_score = sum(1 for kw in self.HEADER_KEYWORDS if kw in row_str)
            # 非空列数越多越像表头
            fill_score = len(non_null) / max(len(row), 1)
            # 文本占比越高越像表头
            text_score = (len(non_null) - num_count) / max(len(non_null), 1)

            score = kw_score * 10 + fill_score * 3 + text_score * 2

            if score > best_score:
                best_score = score
                best_row = i

        return best_row

    def _is_numeric(self, val) -> bool:
        if isinstance(val, (int, float)):
            return True
        s = str(val).strip().replace(",", "").replace("，", "")
        s = re.sub(r"[元¥￥$]", "", s)
        try:
            float(s)
            return True
        except (ValueError, TypeError):
            return False

    def _extract_table_from_header(self, raw_df: pd.DataFrame, header_row: int) -> pd.DataFrame:
        """从指定行作为列头提取数据（单行表头）"""
        if header_row >= len(raw_df):
            return pd.DataFrame()

        # 仅用 header_row 一行作为列名，避免数据被误识别为二级表头
        headers = []
        for col_idx in range(len(raw_df.columns)):
            h = raw_df.iloc[header_row, col_idx]
            h_str = str(h).strip() if pd.notna(h) else ""
            headers.append(h_str if h_str else f"列{col_idx}")

        # 处理列名重复
        seen = {}
        unique_headers = []
        for h in headers:
            if h in seen:
                seen[h] += 1
                unique_headers.append(f"{h}_{seen[h]}")
            else:
                seen[h] = 0
                unique_headers.append(h)

        # 数据从 header_row + 1 开始
        data = raw_df.iloc[header_row + 1:].reset_index(drop=True)
        data.columns = unique_headers

        # 清理空列
        data = data.dropna(axis=1, how="all")

        # 对字符串列做 strip，防止前后空格导致SQL匹配失败
        for col in data.columns:
            if data[col].dtype == object:
                data[col] = data[col].apply(
                    lambda x: x.strip() if isinstance(x, str) else x
                )

        return data

    def _parse_csv(self, path: Path, result: ParseResult):
        for encoding in ["utf-8", "gbk", "gb2312", "utf-8-sig"]:
            try:
                df = pd.read_csv(str(path), encoding=encoding)
                break
            except (UnicodeDecodeError, pd.errors.ParserError):
                continue
        else:
            raise ValueError(f"无法解析CSV文件编码: {path.name}")

        df = df.dropna(how="all").reset_index(drop=True)
        if not df.empty:
            result.tables.append(TableBlock(
                dataframe=df,
                source_file=path.name,
                metadata={"total_rows": len(df), "columns": list(df.columns)},
            ))

    # ---------- 工具 ----------
    def _clean_table_headers(self, df: pd.DataFrame) -> pd.DataFrame:
        first_row = df.iloc[0]
        if all(isinstance(v, str) and not re.match(r"^[\d.,]+$", v) for v in first_row if pd.notna(v)):
            df.columns = [str(v).strip() if pd.notna(v) else f"列{i}" for i, v in enumerate(first_row)]
            df = df.iloc[1:].reset_index(drop=True)
        return df
